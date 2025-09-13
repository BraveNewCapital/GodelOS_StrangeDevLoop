"""
Knowledge Ingestion Service

Comprehensive service for ingesting knowledge from various sources including
URLs, files, Wikipedia, and manual text input with robust processing and validation.
"""

import asyncio
from asyncio import TimeoutError as AsyncioTimeoutError

# Optimized timeout for fast processing - aggressive fallback to basic processing
CONTENT_PROCESS_TIMEOUT = 60  # Reduced to 1 minute - fail fast and use fallback
import hashlib
import json
import logging
import os
import tempfile
import time
import traceback
import uuid
from typing import Dict, List, Optional, Any, Tuple
from pathlib import Path

import aiofiles

# Optional imports for file processing
try:
    import PyPDF2
    HAS_PDF = True
except ImportError:
    HAS_PDF = False
    PyPDF2 = None

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    Document = None

from .knowledge_models import (
    ImportRequest, URLImportRequest, FileImportRequest, WikipediaImportRequest,
    TextImportRequest, BatchImportRequest, ImportProgress, ContentChunk,
    KnowledgeItem, ImportSource, ImportStatistics
)
from .external_apis import wikipedia_api, web_scraper, content_processor
from .persistence import get_persistence_layer

# Will be set by main.py to avoid circular imports
knowledge_management_service = None

# Import knowledge pipeline service
from .knowledge_pipeline_service import knowledge_pipeline_service

logger = logging.getLogger(__name__)


async def extract_text_from_pdf(file_path: str) -> str:
    """Extract text content from PDF file using PyPDF2."""
    if not HAS_PDF:
        raise ValueError("PDF processing not available - install PyPDF2")
    
    try:
        text_content = []
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            
            # Extract text from each page
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text.strip():
                        text_content.append(f"--- Page {page_num + 1} ---\n{page_text.strip()}")
                except Exception as e:
                    logger.warning(f"Could not extract text from page {page_num + 1}: {e}")
                    text_content.append(f"--- Page {page_num + 1} ---\n[Text extraction failed]")
        
        if not text_content:
            return "No readable text found in PDF"
        
        full_text = "\n\n".join(text_content)
        logger.info(f"Successfully extracted {len(full_text)} characters from PDF with {len(pdf_reader.pages)} pages")
        return full_text
        
    except Exception as e:
        logger.error(f"Error extracting text from PDF {file_path}: {e}")
        raise ValueError(f"Failed to extract text from PDF: {str(e)}")


async def extract_text_from_docx(file_path: str) -> str:
    """Extract text content from DOCX file using python-docx."""
    if not HAS_DOCX:
        raise ValueError("DOCX processing not available - install python-docx")
    
    try:
        doc = Document(file_path)
        text_content = []
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text.strip())
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_content.append(" | ".join(row_text))
        
        full_text = "\n\n".join(text_content)
        logger.info(f"Successfully extracted {len(full_text)} characters from DOCX")
        return full_text if full_text else "No readable text found in DOCX"
        
    except Exception as e:
        logger.error(f"Error extracting text from DOCX {file_path}: {e}")
        raise ValueError(f"Failed to extract text from DOCX: {str(e)}")


class KnowledgeIngestionService:
    """Main service for knowledge ingestion operations."""
    
    def __init__(self, storage_path: str = "./knowledge_storage"):
        self.storage_path = Path(storage_path)
        self.storage_path.mkdir(exist_ok=True)
        self.active_imports: Dict[str, ImportProgress] = {}
        self.knowledge_store: Dict[str, KnowledgeItem] = {}
        self.import_queue: asyncio.Queue = asyncio.Queue()
        self.processing_task: Optional[asyncio.Task] = None
        self.max_concurrent_imports = 5
        self.semaphore = asyncio.Semaphore(self.max_concurrent_imports)
        self.websocket_manager = None  # Will be set from main.py
        logger.info(f"🔍 INIT: KnowledgeIngestionService created, websocket_manager: {self.websocket_manager}")
    
    async def initialize(self, websocket_manager=None):
        """Initialize the knowledge ingestion service."""
        logger.info("Initializing Knowledge Ingestion Service...")
        
        # Set websocket manager if provided
        if websocket_manager:
            self.websocket_manager = websocket_manager
            logger.info(f"🔍 INITIALIZE: WebSocket manager set during initialization: {self.websocket_manager is not None}")
        
        # Initialize external APIs
        await wikipedia_api.initialize()
        await web_scraper.initialize()
        
        # Start processing task
        self.processing_task = asyncio.create_task(self._process_import_queue())
        
        # Load existing knowledge items
        await self._load_existing_knowledge()
        
        logger.info("Knowledge Ingestion Service initialized successfully")
    
    async def shutdown(self):
        """Shutdown the knowledge ingestion service."""
        logger.info("Shutting down Knowledge Ingestion Service...")
        
        # Cancel processing task
        if self.processing_task:
            self.processing_task.cancel()
            try:
                await self.processing_task
            except asyncio.CancelledError:
                pass
        
        # Shutdown external APIs
        await wikipedia_api.shutdown()
        await web_scraper.shutdown()
        
        logger.info("Knowledge Ingestion Service shutdown complete")
    
    async def import_from_url(self, request: URLImportRequest) -> str:
        """Import knowledge from a URL."""
        import_id = str(uuid.uuid4())
        
        progress = ImportProgress(
            import_id=import_id,
            status="queued",
            progress_percentage=0.0,
            current_step="Queued for processing",
            total_steps=5,
            completed_steps=0,
            started_at=time.time()
        )
        
        self.active_imports[import_id] = progress
        
        # Add to processing queue
        await self.import_queue.put(("url", import_id, request))
        
        logger.info(f"URL import queued: {import_id} for {request.url}")
        return import_id
    
    async def import_from_file(self, request: FileImportRequest, file_content: bytes) -> str:
        """Import knowledge from an uploaded file."""
        import_id = str(uuid.uuid4())
        
        progress = ImportProgress(
            import_id=import_id,
            status="queued",
            progress_percentage=0.0,
            current_step="Queued for processing",
            total_steps=6,
            completed_steps=0,
            started_at=time.time()
        )
        
        self.active_imports[import_id] = progress
        
        # Save file temporarily
        # Ensure we do not accidentally preserve any client-side path components
        # that may be present in the uploaded filename (e.g., "tmp/test_upload.txt").
        safe_name = Path(request.filename).name
        temp_file_path = self.storage_path / f"temp_{import_id}_{safe_name}"
        async with aiofiles.open(temp_file_path, 'wb') as f:
            await f.write(file_content)
        
        # Add to processing queue
        await self.import_queue.put(("file", import_id, request, str(temp_file_path)))
        
        logger.info(f"File import queued: {import_id} for {request.filename}")
        return import_id
    
    async def import_from_wikipedia(self, request: WikipediaImportRequest) -> str:
        """Import knowledge from Wikipedia."""
        import_id = str(uuid.uuid4())
        
        progress = ImportProgress(
            import_id=import_id,
            status="queued",
            progress_percentage=0.0,
            current_step="Queued for processing",
            total_steps=4,
            completed_steps=0,
            started_at=time.time()
        )
        
        self.active_imports[import_id] = progress
        
        # Add to processing queue
        await self.import_queue.put(("wikipedia", import_id, request))
        
        # IMMEDIATE FIX: Start background task to send progress updates
        logger.info(f"🔍 TASK: Creating background progress task for {import_id}")
        asyncio.create_task(self._send_wikipedia_progress_updates(import_id, request))
        logger.info(f"🔍 TASK: Background task created for {import_id}")
        
        logger.info(f"Wikipedia import queued: {import_id} for {request.page_title}")
        return import_id
    
    async def import_from_text(self, request: TextImportRequest) -> str:
        """Import knowledge from manual text input."""
        import_id = str(uuid.uuid4())
        
        progress = ImportProgress(
            import_id=import_id,
            status="queued",
            progress_percentage=0.0,
            current_step="Queued for processing",
            total_steps=3,
            completed_steps=0,
            started_at=time.time()
        )
        
        self.active_imports[import_id] = progress
        
        # Add to processing queue
        await self.import_queue.put(("text", import_id, request))
        
        logger.info(f"Text import queued: {import_id}")
        return import_id
    
    async def batch_import(self, request: BatchImportRequest) -> List[str]:
        """Process multiple import requests in batch."""
        import_ids = []
        
        for import_request in request.import_requests:
            if isinstance(import_request, URLImportRequest):
                import_id = await self.import_from_url(import_request)
            elif isinstance(import_request, WikipediaImportRequest):
                import_id = await self.import_from_wikipedia(import_request)
            elif isinstance(import_request, TextImportRequest):
                import_id = await self.import_from_text(import_request)
            else:
                logger.warning(f"Unsupported import request type: {type(import_request)}")
                continue
            
            import_ids.append(import_id)
        
        logger.info(f"Batch import queued: {len(import_ids)} items")
        return import_ids
    
    async def get_import_progress(self, import_id: str) -> Optional[ImportProgress]:
        """Get the progress of an import operation from memory or persistence."""
        # Check memory first
        if import_id in self.active_imports:
            return self.active_imports[import_id]
        
        # Check persistence if not in memory
        try:
            persistence = await get_persistence_layer()
            progress_data = await persistence.import_tracker.load_progress(import_id)
            if progress_data:
                # Reconstruct ImportProgress object
                progress = ImportProgress(
                    import_id=progress_data["import_id"],
                    status=progress_data["status"],
                    progress_percentage=progress_data["progress_percentage"],
                    current_step=progress_data["current_step"],
                    total_steps=progress_data["total_steps"],
                    completed_steps=progress_data["completed_steps"],
                    started_at=progress_data["started_at"],
                    estimated_completion=progress_data.get("estimated_completion"),
                    error_message=progress_data.get("error_message"),
                    warnings=progress_data.get("warnings", [])
                )
                # Add back to memory cache
                self.active_imports[import_id] = progress
                return progress
        except Exception as e:
            logger.error(f"Error loading import progress from persistence: {e}")
        
        return None
    
    async def cancel_import(self, import_id: str) -> bool:
        """Cancel an import operation."""
        if import_id in self.active_imports:
            progress = self.active_imports[import_id]
            if progress.status in ["queued", "processing"]:
                progress.status = "cancelled"
                progress.error_message = "Import cancelled by user"
                await self._broadcast_progress_update(import_id, progress)
                await self._broadcast_completion(import_id, False, "Import cancelled")
                logger.info(f"Import cancelled: {import_id}")
                return True
        return False
    
    async def reset_stuck_imports(self) -> int:
        """Reset any stuck imports that have been processing too long."""
        current_time = time.time()
        stuck_imports = []
        
        for import_id, progress in self.active_imports.items():
            if progress.status == "processing":
                # Check if import has been processing for more than 15 minutes
                if current_time - progress.started_at > 900:  # 15 minutes
                    stuck_imports.append(import_id)
        
        for import_id in stuck_imports:
            logger.warning(f"🔄 RESET: Resetting stuck import {import_id}")
            await self.cancel_import(import_id)
        
        return len(stuck_imports)
    
    async def _broadcast_progress_update(self, import_id: str, progress: ImportProgress):
        """Broadcast progress update via WebSocket and save to persistence."""
        logger.info(f"🔍 DEBUG: _broadcast_progress_update called for {import_id}")
        
        # Save progress to persistence
        try:
            persistence = await get_persistence_layer()
            progress_data = {
                "import_id": import_id,
                "status": progress.status,
                "progress_percentage": progress.progress_percentage,
                "current_step": progress.current_step,
                "total_steps": progress.total_steps,
                "completed_steps": progress.completed_steps,
                "started_at": progress.started_at,
                "estimated_completion": progress.estimated_completion,
                "error_message": progress.error_message,
                "warnings": progress.warnings
            }
            await persistence.import_tracker.store_progress(import_id, progress_data)
            logger.debug(f"Saved import progress for {import_id} to persistence")
        except Exception as e:
            logger.error(f"Error saving import progress to persistence: {e}")
        
        # Broadcast via WebSocket if available
        logger.info(f"🔍 DEBUG: websocket_manager exists: {self.websocket_manager is not None}")
        
        if self.websocket_manager:
            has_connections = self.websocket_manager.has_connections()
            logger.info(f"🔍 DEBUG: has_connections: {has_connections}")
            
            if has_connections:
                try:
                    progress_event = {
                        "type": "import_progress",
                        "timestamp": time.time(),
                        "import_id": import_id,
                        "progress": progress.progress_percentage,
                        "status": progress.status,
                        "message": progress.current_step,
                        "completed_steps": progress.completed_steps,
                        "total_steps": progress.total_steps
                    }
                    logger.info(f"🔍 DEBUG: Broadcasting progress event: {progress_event}")
                    await self.websocket_manager.broadcast(progress_event)
                    logger.info(f"🔍 DEBUG: Successfully broadcasted progress update for import {import_id}")
                except Exception as e:
                    logger.error(f"🔍 DEBUG: Failed to broadcast progress update: {e}")
                    import traceback
                    logger.error(f"🔍 DEBUG: Traceback: {traceback.format_exc()}")
            else:
                logger.warning(f"🔍 DEBUG: No WebSocket connections available for progress update")
        else:
            logger.warning(f"🔍 DEBUG: websocket_manager is None, cannot broadcast progress")
    
    async def _broadcast_completion(self, import_id: str, success: bool, message: str = None):
        """Broadcast import completion via WebSocket if available."""
        logger.info(f"🔍 DEBUG: _broadcast_completion called for {import_id}, success: {success}")
        
        if self.websocket_manager and self.websocket_manager.has_connections():
            try:
                event_type = "import_completed" if success else "import_failed"
                completion_event = {
                    "type": event_type,
                    "timestamp": time.time(),
                    "import_id": import_id,
                    "message": message or ("Import completed successfully" if success else "Import failed"),
                    "success": success
                }
                logger.info(f"🔍 DEBUG: Broadcasting completion event: {completion_event}")
                await self.websocket_manager.broadcast(completion_event)
                logger.info(f"🔍 DEBUG: Successfully broadcasted import completion for {import_id}: {event_type}")
            except Exception as e:
                logger.error(f"🔍 DEBUG: Failed to broadcast completion event: {e}")
                import traceback
                logger.error(f"🔍 DEBUG: Traceback: {traceback.format_exc()}")
        else:
            logger.warning(f"🔍 DEBUG: Cannot broadcast completion - websocket_manager or connections not available")
    
    async def _process_import_queue(self):
        """Background task to process the import queue."""
        logger.info("🔍 DEBUG: _process_import_queue started")
        while True:
            try:
                logger.info(f"🔍 DEBUG: Waiting for import from queue, current size: {self.import_queue.qsize()}")
                # Get next import from queue
                import_data = await self.import_queue.get()
                logger.info(f"🔍 DEBUG: Got import from queue: {import_data[0]}")
                
                # Process with semaphore to limit concurrent imports
                async with self.semaphore:
                    await self._process_single_import(import_data)
                
                self.import_queue.task_done()
                logger.info(f"🔍 DEBUG: Import processing completed for {import_data[0]}")
                
            except asyncio.CancelledError:
                logger.info("🔍 DEBUG: _process_import_queue cancelled")
                break
            except Exception as e:
                logger.error(f"Error in import queue processing: {e}")
                await asyncio.sleep(1)
    
    async def _process_single_import(self, import_data: Tuple):
        """Process a single import operation with comprehensive error handling."""
        import_type = import_data[0]
        import_id = import_data[1]
        request = import_data[2]
        
        # Add overall timeout for the entire import process
        IMPORT_TIMEOUT = 600  # 10 minutes max per import
        
        try:
            progress = self.active_imports[import_id]
            progress.status = "processing"
            progress.current_step = "Starting import"
            
            logger.info(f"🔍 DEBUG: Starting processing for import {import_id}, type: {import_type}")
            
            # Broadcast initial progress update
            await self._broadcast_progress_update(import_id, progress)
            logger.info(f"🔍 DEBUG: Broadcasted initial progress for {import_id}")
            
            # Wrap the actual processing with a timeout to prevent stuck imports
            try:
                if import_type == "url":
                    logger.info(f"🔍 DEBUG: Processing URL import {import_id}")
                    await asyncio.wait_for(
                        self._process_url_import(import_id, request),
                        timeout=IMPORT_TIMEOUT
                    )
                elif import_type == "file":
                    logger.info(f"🔍 DEBUG: Processing file import {import_id}")
                    file_path = import_data[3]
                    await asyncio.wait_for(
                        self._process_file_import(import_id, request, file_path),
                        timeout=IMPORT_TIMEOUT
                    )
                elif import_type == "wikipedia":
                    logger.info(f"🔍 DEBUG: Processing Wikipedia import {import_id}")
                    await asyncio.wait_for(
                        self._process_wikipedia_import(import_id, request),
                        timeout=IMPORT_TIMEOUT
                    )
                elif import_type == "text":
                    logger.info(f"🔍 DEBUG: Processing text import {import_id}")
                    await asyncio.wait_for(
                        self._process_text_import(import_id, request),
                        timeout=IMPORT_TIMEOUT
                    )
                    
            except asyncio.TimeoutError:
                logger.error(f"❌ TIMEOUT: Import {import_id} timed out after {IMPORT_TIMEOUT} seconds")
                progress.status = "failed"
                progress.error_message = f"Import timed out after {IMPORT_TIMEOUT} seconds"
                await self._broadcast_progress_update(import_id, progress)
                await self._broadcast_completion(import_id, False, progress.error_message)
                return
            
            logger.info(f"🔍 DEBUG: Individual processing completed for {import_id}")
            
            # Mark as completed
            progress.status = "completed"
            progress.progress_percentage = 100.0
            progress.current_step = "Import completed successfully"
            progress.completed_steps = progress.total_steps
            
            # Broadcast final progress update
            await self._broadcast_progress_update(import_id, progress)
            logger.info(f"🔍 DEBUG: Broadcasted final progress for {import_id}")
            
            # Broadcast completion event
            await self._broadcast_completion(import_id, True, "Import completed successfully")
            logger.info(f"🔍 DEBUG: Broadcasted completion event for {import_id}")
            
            logger.info(f"Import completed successfully: {import_id}")
            
        except Exception as e:
            logger.error(f"Import failed: {import_id} - {e}")
            logger.error(f"🔍 DEBUG: Exception details: {type(e).__name__}: {str(e)}")
            import traceback
            logger.error(f"🔍 DEBUG: Traceback: {traceback.format_exc()}")
            
            progress = self.active_imports[import_id]
            progress.status = "failed"
            progress.error_message = str(e)
            
            # Broadcast failure
            await self._broadcast_progress_update(import_id, progress)
            await self._broadcast_completion(import_id, False, str(e))
    
    async def _process_content(self, content: str, title: str, metadata: Dict[str, Any]) -> Dict[str, Any]:
        """Process raw content using basic knowledge extraction (temporarily bypassing advanced pipeline)."""
        try:
            logger.info(f"🔍 DEBUG: _process_content called with title: {title}")
            
            # TEMPORARILY USE BASIC PROCESSING ONLY
            logger.info("🔄 Using basic knowledge extraction processing")
            
            # Basic processing for backward compatibility
            cleaned_content = content_processor.clean_text(content)
            sentences = content_processor.extract_sentences(cleaned_content)
            chunks = content_processor.chunk_content(cleaned_content)
            keywords = content_processor.extract_keywords(cleaned_content)
            language = content_processor.detect_language(cleaned_content)
            
            result = {
                'title': title,
                'content': cleaned_content,
                'sentences': sentences,
                'chunks': chunks,
                'keywords': keywords,
                'language': language,
                'metadata': metadata,
                'word_count': len(cleaned_content.split()),
                'char_count': len(cleaned_content),
                'entities_extracted': 0,
                'relationships_extracted': 0,
                'knowledge_items': []
            }
            logger.info(f"🔍 DEBUG: _process_content returning with {len(result)} keys")
            return result
        except Exception as e:
            logger.error(f"❌ Error in content processing: {e}")
            logger.error(f"🔍 DEBUG: Exception traceback: {traceback.format_exc()}")
            # Fallback to basic processing on error
            cleaned_content = content_processor.clean_text(content)
            return {
                'title': title,
                'content': cleaned_content,
                'metadata': metadata,
                'word_count': len(cleaned_content.split()),
                'char_count': len(cleaned_content),
                'processing_error': str(e)
            }
    
    async def _load_existing_knowledge(self):
        """Load existing knowledge items from storage."""
        try:
            for item_file in self.storage_path.glob("*.json"):
                # Skip temp files and category listings
                if item_file.name.startswith("temp_") or item_file.name == "categories.json":
                    continue
                
                try:
                    async with aiofiles.open(item_file, 'r') as f:
                        item_data = json.loads(await f.read())
                        # Guard: some legacy files may contain a list; skip invalid shapes
                        if isinstance(item_data, list):
                            logger.warning(f"Skipping non-mapping knowledge file {item_file} (list detected)")
                            continue
                        if not isinstance(item_data, dict):
                            logger.warning(f"Skipping invalid knowledge file {item_file} (type={type(item_data)})")
                            continue
                        knowledge_item = KnowledgeItem(**item_data)
                        self.knowledge_store[knowledge_item.id] = knowledge_item
                except Exception as e:
                    logger.warning(f"Failed to load knowledge item {item_file}: {e}")
            
            logger.info(f"Loaded {len(self.knowledge_store)} existing knowledge items")
        except Exception as e:
            logger.error(f"Error loading existing knowledge: {e}")

    async def _process_text_import(self, import_id: str, request):
        """Process text import request."""
        try:
            progress = self.active_imports[import_id]
            progress.current_step = "Processing text content"
            progress.progress_percentage = 25.0
            progress.completed_steps = 1
            
            # Broadcast progress update
            await self._broadcast_progress_update(import_id, progress)
            
            # Extract content and metadata
            content = request.content
            title = request.title or "Manual Text Entry"
            
            # Process the content
            progress.current_step = "Analyzing content"
            progress.progress_percentage = 50.0
            progress.completed_steps = 2
            
            # Broadcast progress update
            await self._broadcast_progress_update(import_id, progress)
            
            try:
                logger.info(f"🔍 DEBUG: Starting _process_content for file import {import_id} with timeout {CONTENT_PROCESS_TIMEOUT}s")
                processed_data = await asyncio.wait_for(
                    self._process_content(
                        content=content,
                        title=title,
                        metadata=request.source.metadata
                    ),
                    timeout=CONTENT_PROCESS_TIMEOUT
                )
                logger.info(f"🔍 DEBUG: _process_content completed for file import {import_id}")
            except AsyncioTimeoutError:
                logger.error(f"❌ Timeout during content processing for file import {import_id}")
                progress.status = "failed"
                progress.error_message = "Content processing timed out"
                await self._broadcast_progress_update(import_id, progress)
                await self._broadcast_completion(import_id, False, progress.error_message)
                return
            except Exception as e:
                logger.error(f"❌ Exception during content processing for file import {import_id}: {e}")
                logger.error(f"🔍 DEBUG: Traceback: {traceback.format_exc()}")
                progress.status = "failed"
                progress.error_message = str(e)
                await self._broadcast_progress_update(import_id, progress)
                await self._broadcast_completion(import_id, False, progress.error_message)
                return
            
            # Create knowledge item
            progress.current_step = "Creating knowledge item"
            progress.progress_percentage = 75.0
            
            # Broadcast progress update
            await self._broadcast_progress_update(import_id, progress)
            
            knowledge_item = KnowledgeItem(
                id=f"text-{import_id}",
                content=processed_data['content'],
                knowledge_type="fact",  # Default type, could be enhanced with classification
                title=processed_data['title'],
                source=request.source,
                import_id=import_id,
                confidence=0.9,  # High confidence for manual entry
                quality_score=0.85,  # Good quality score for manual entry
                categories=request.categorization_hints or ["manual"],
                auto_categories=[],
                manual_categories=request.categorization_hints or ["manual"],
                metadata={
                    'word_count': processed_data['word_count'],
                    'char_count': processed_data['char_count'],
                    'chunks': len(processed_data['chunks']),
                    'keywords': processed_data['keywords'],
                    'language': processed_data['language']
                }
            )
            
            # Store the knowledge item
            await self._store_knowledge_item(knowledge_item)
            self.knowledge_store[knowledge_item.id] = knowledge_item
            
            logger.info(f"Text import processed successfully: {import_id}")
            
        except Exception as e:
            logger.error(f"Failed to process text import {import_id}: {e}")
            raise

    async def _process_url_import(self, import_id: str, request):
        """Process URL import request with real web scraping and progress tracking."""
        try:
            progress = self.active_imports[import_id]
            progress.current_step = "Connecting to URL"
            progress.progress_percentage = 15.0
            progress.completed_steps = 1
            
            # Broadcast progress update
            await self._broadcast_progress_update(import_id, progress)
            
            # Get URL from request
            url = str(getattr(request, 'url', getattr(request, 'source_url', 'Unknown')))
            
            progress.current_step = f"Fetching content from {url}"
            progress.progress_percentage = 30.0
            progress.completed_steps = 2
            await self._broadcast_progress_update(import_id, progress)
            
            # Use the real web scraper
            logger.info(f"🔍 DEBUG: Scraping URL: {url}")
            scraped_data = await web_scraper.scrape_url(url)
            
            if scraped_data.get('is_fallback'):
                logger.warning(f"Using fallback content for {url}")
            
            content = scraped_data.get('content', '')
            title = scraped_data.get('title', f"Web Content from {url}")
            
            # Update progress with content size information
            word_count = scraped_data.get('word_count', len(content.split()))
            progress.current_step = f"Processing {word_count} words from web page"
            progress.progress_percentage = 55.0
            progress.completed_steps = 3
            
            # Broadcast progress update
            await self._broadcast_progress_update(import_id, progress)
            
            # Process the actual content
            logger.info(f"🔍 DEBUG: Processing scraped content for {import_id}")
            metadata = {
                'source_url': url,
                'description': scraped_data.get('description', ''),
                'keywords': scraped_data.get('metadata', {}).get('keywords', []),
                'word_count': word_count,
                'char_count': scraped_data.get('char_count', len(content)),
                'content_type': scraped_data.get('metadata', {}).get('content_type', ''),
                'is_fallback': scraped_data.get('is_fallback', False)
            }
            
            # Merge with any existing metadata from request
            if hasattr(request, 'source') and hasattr(request.source, 'metadata'):
                metadata.update(request.source.metadata)
            
            progress.current_step = "Processing and extracting knowledge"
            progress.progress_percentage = 75.0
            progress.completed_steps = 4
            await self._broadcast_progress_update(import_id, progress)
            
            try:
                logger.info(f"🔍 DEBUG: Starting _process_content for URL import {import_id} with timeout {CONTENT_PROCESS_TIMEOUT}s")
                processed_data = await asyncio.wait_for(
                    self._process_content(
                        content=content,
                        title=title,
                        metadata=metadata
                    ),
                    timeout=CONTENT_PROCESS_TIMEOUT
                )
                logger.info(f"🔍 DEBUG: _process_content completed for URL import {import_id}")
            except AsyncioTimeoutError:
                logger.error(f"❌ Timeout during content processing for URL import {import_id}")
                progress.status = "failed"
                progress.error_message = "Content processing timed out"
                await self._broadcast_progress_update(import_id, progress)
                await self._broadcast_completion(import_id, False, progress.error_message)
                return
            except Exception as e:
                logger.error(f"❌ Exception during content processing for URL import {import_id}: {e}")
                logger.error(f"🔍 DEBUG: Traceback: {traceback.format_exc()}")
                progress.status = "failed"
                progress.error_message = str(e)
                await self._broadcast_progress_update(import_id, progress)
                await self._broadcast_completion(import_id, False, progress.error_message)
                return
            
            progress.current_step = "Creating knowledge item"
            progress.progress_percentage = 90.0
            progress.completed_steps = 5
            
            # Broadcast progress update
            await self._broadcast_progress_update(import_id, progress)
            
            # Create knowledge item with real data
            knowledge_item = KnowledgeItem(
                id=f"url-{import_id}",
                content=processed_data['content'],
                knowledge_type="fact",
                title=processed_data['title'],
                source=ImportSource(
                    source_type="url",
                    source_identifier=url,
                    metadata=metadata
                ),
                import_id=import_id,
                confidence=0.8 if not scraped_data.get('is_fallback') else 0.4,
                quality_score=0.75,
                categories=request.categorization_hints or ["web"],
                auto_categories=[],
                manual_categories=request.categorization_hints or ["web"],
                relationships=[],
                metadata={
                    **processed_data.get('metadata', {}),
                    'source_url': request.url,
                    'max_depth': request.max_depth
                }
            )
            
            # Store the knowledge item
            await self._store_knowledge_item(knowledge_item)
            self.knowledge_store[knowledge_item.id] = knowledge_item
            
            logger.info(f"URL import processed successfully: {import_id}")
            
        except Exception as e:
            logger.error(f"Failed to process URL import {import_id}: {e}")
            raise

    async def _process_file_import(self, import_id: str, request, file_path: str):
        """Process file import request."""
        try:
            progress = self.active_imports[import_id]
            progress.current_step = "Reading file content"
            progress.progress_percentage = 25.0
            progress.completed_steps = 1
            
            # Broadcast progress update
            await self._broadcast_progress_update(import_id, progress)
            
            # Read file content based on type with proper extraction
            content = ""
            enhanced_content_data = None
            
            if request.file_type == "pdf":
                if not HAS_PDF:
                    raise ValueError("PDF processing not available - install PyPDF2")
                logger.info(f"Extracting text from PDF: {file_path}")
                raw_content = await extract_text_from_pdf(file_path)
                logger.info(f"🔍 PDF DEBUG: Extracted {len(raw_content)} characters from PDF")
                logger.info(f"🔍 PDF DEBUG: First 200 chars: {repr(raw_content[:200])}")
                logger.info(f"🔍 PDF DEBUG: Content preview: {raw_content[:500] if raw_content else 'EMPTY'}")
                
                # Apply aggressive size limits for efficiency
                MAX_PDF_CONTENT = 75000  # 75K character limit for PDF processing
                if len(raw_content) > MAX_PDF_CONTENT:
                    logger.warning(f"🔍 PDF OPTIMIZATION: Large PDF content ({len(raw_content)} chars), truncating to {MAX_PDF_CONTENT} for efficiency")
                    raw_content = raw_content[:MAX_PDF_CONTENT]
                
                # Use the existing knowledge pipeline service for semantic analysis
                logger.info(f"🔍 PDF ENHANCED: Processing PDF content with advanced knowledge pipeline")
                logger.info(f"🔍 PDF DEBUG: Pipeline service available: {knowledge_pipeline_service is not None}")
                logger.info(f"🔍 PDF DEBUG: Pipeline service initialized: {knowledge_pipeline_service.initialized if knowledge_pipeline_service else False}")
                
                try:
                    # Use the existing knowledge pipeline service that has spaCy and HuggingFace models
                    if knowledge_pipeline_service and knowledge_pipeline_service.initialized:
                        logger.info(f"🔍 PDF DEBUG: Processing {len(raw_content)} characters through pipeline with timeout {CONTENT_PROCESS_TIMEOUT}s")
                        try:
                            pipeline_result = await asyncio.wait_for(
                                knowledge_pipeline_service.process_text_document(
                                    content=raw_content,
                                    title=request.filename,
                                    metadata={
                                        'file_type': request.file_type,
                                        'filename': request.filename,
                                        'encoding': request.encoding,
                                        'source': 'pdf_extraction'
                                    }
                                ),
                                timeout=CONTENT_PROCESS_TIMEOUT
                            )
                        except asyncio.TimeoutError:
                            logger.error(f"❌ Timeout during PDF pipeline processing for import {import_id}")
                            # Continue with basic processing instead of failing
                            pipeline_result = None
                        
                        logger.info(f"🔍 PDF DEBUG: Pipeline result keys: {list(pipeline_result.keys()) if pipeline_result else 'None'}")
                        logger.info(f"🔍 PDF DEBUG: Pipeline result entities count: {pipeline_result.get('entities_extracted', 0)}")
                        logger.info(f"🔍 PDF DEBUG: Pipeline result relationships count: {pipeline_result.get('relationships_extracted', 0)}")
                        logger.info(f"🔍 PDF DEBUG: Pipeline result knowledge items: {len(pipeline_result.get('knowledge_items', []))}")
                        
                        # Extract semantic concepts from the pipeline results using the correct keys
                        entities_count = pipeline_result.get('entities_extracted', 0)
                        relationships_count = pipeline_result.get('relationships_extracted', 0)
                        knowledge_items = pipeline_result.get('knowledge_items', [])
                        
                        # CRITICAL FIX: Get the actual extracted entities from the pipeline result
                        # The pipeline stores extracted data in the 'processed_data' key
                        processed_pipeline_data = pipeline_result.get('processed_data', {})
                        raw_entities = processed_pipeline_data.get('entities', [])
                        raw_relationships = processed_pipeline_data.get('relationships', [])
                        
                        logger.info(f"🔍 PDF DEBUG: Raw entities from pipeline: {len(raw_entities)}")
                        logger.info(f"🔍 PDF DEBUG: First 3 raw entities: {raw_entities[:3] if raw_entities else 'NONE'}")
                        logger.info(f"🔍 PDF DEBUG: Raw relationships from pipeline: {len(raw_relationships)}")
                        
                        # Extract meaningful entity names from the ACTUAL pipeline results with SMART FILTERING
                        meaningful_entities = []
                        meaningful_relationships = []
                        semantic_concepts = []
                        
                        # Define filtering criteria for meaningful concepts
                        def is_meaningful_concept(text: str, entity_label: str = None) -> bool:
                            """Filter for semantically meaningful concepts only."""
                            if not text or len(text.strip()) < 3:
                                return False
                            
                            text = text.strip()
                            
                            # Skip generic/noise terms
                            noise_terms = {
                                'file', 'document', 'pdf', 'docx', 'txt', 'upload', 'download',
                                'the', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for', 'of', 'with',
                                'this', 'that', 'these', 'those', 'what', 'where', 'when', 'how', 'why',
                                'today', 'yesterday', 'tomorrow', 'now', 'then', 'here', 'there'
                            }
                            if text.lower() in noise_terms:
                                return False
                            
                            # Skip pure numbers or single characters
                            if text.isdigit() or len(text) == 1:
                                return False
                            
                            # Skip file extensions and paths
                            if '.' in text and len(text.split('.')[-1]) <= 4:
                                return False
                            
                            # Prioritize meaningful entity types
                            if entity_label:
                                high_value_labels = {'ORG', 'PERSON', 'PRODUCT', 'TECHNOLOGY', 'SYSTEM'}
                                low_value_labels = {'CARDINAL', 'ORDINAL', 'DATE', 'TIME', 'GPE'}
                                
                                if entity_label in high_value_labels:
                                    return True  # Always include organizations, people, products
                                elif entity_label in low_value_labels:
                                    # Only include if it's a substantial term
                                    return len(text) > 4 and not text.isdigit()
                            
                            # Include multi-word technical terms (likely meaningful)
                            if ' ' in text and len(text) > 6:
                                return True
                            
                            # Include capitalized terms (likely proper nouns)
                            if text[0].isupper() and len(text) > 3:
                                return True
                            
                            # Include terms with mixed case (likely technical/compound terms)
                            if any(c.isupper() for c in text[1:]) and len(text) > 4:
                                return True
                            
                            return False
                        
                        if raw_entities:
                            logger.info(f"🔍 FILTERING: Processing {len(raw_entities)} raw entities")
                            for i, entity in enumerate(raw_entities):
                                try:
                                    if isinstance(entity, dict) and 'text' in entity:
                                        entity_text = entity['text'].strip()
                                        entity_label = entity.get('label', '')
                                        
                                        if is_meaningful_concept(entity_text, entity_label):
                                            meaningful_entities.append(entity_text)
                                            semantic_concepts.append(entity_text)
                                            logger.info(f"🔍 FILTERED: Kept meaningful entity '{entity_text}' ({entity_label})")
                                        else:
                                            logger.info(f"🔍 FILTERED: Skipped noise entity '{entity_text}' ({entity_label})")
                                    
                                    # Limit entity processing to prevent timeout
                                    if len(semantic_concepts) >= 20:  # Cap at 20 meaningful entities
                                        logger.info(f"🔍 FILTERING: Reached entity limit (20), stopping entity processing")
                                        break
                                        
                                except Exception as e:
                                    logger.warning(f"🔍 FILTERING: Error processing entity {i}: {e}")
                                    continue
                        
                        if raw_relationships:
                            logger.info(f"🔍 FILTERING: Processing {len(raw_relationships)} raw relationships")
                            processed_relationships = 0
                            for i, rel in enumerate(raw_relationships):
                                try:
                                    if isinstance(rel, dict):
                                        source_text = rel.get('source', {}).get('text', '').strip()
                                        target_text = rel.get('target', {}).get('text', '').strip()
                                        source_label = rel.get('source', {}).get('label', '')
                                        target_label = rel.get('target', {}).get('label', '')
                                        relation = rel.get('relation', '').strip()
                                        
                                        # Only create relationships between meaningful concepts
                                        if (source_text and target_text and relation and 
                                            is_meaningful_concept(source_text, source_label) and 
                                            is_meaningful_concept(target_text, target_label)):
                                            rel_description = f"{source_text} {relation} {target_text}"
                                            meaningful_relationships.append(rel_description)
                                            
                                            # Add both entities as concepts if not already present
                                            if source_text not in semantic_concepts:
                                                semantic_concepts.append(source_text)
                                            if target_text not in semantic_concepts:
                                                semantic_concepts.append(target_text)
                                            logger.info(f"🔍 FILTERED: Kept meaningful relationship '{rel_description}'")
                                            processed_relationships += 1
                                        else:
                                            logger.info(f"🔍 FILTERED: Skipped noise relationship '{source_text}' → '{target_text}'")
                                    
                                    # Limit relationship processing to prevent timeout
                                    if processed_relationships >= 15:  # Cap at 15 meaningful relationships
                                        logger.info(f"🔍 FILTERING: Reached relationship limit (15), stopping relationship processing")
                                        break
                                        
                                except Exception as e:
                                    logger.warning(f"🔍 FILTERING: Error processing relationship {i}: {e}")
                                    continue
                        
                        enhanced_metadata = {
                            'pipeline_entities': entities_count,
                            'pipeline_relationships': relationships_count,
                            'pipeline_processing_time': pipeline_result.get('processing_time_seconds', 0),
                            'semantic_concepts': semantic_concepts,  # Real entity names like "Psychometric Engine"
                            'extracted_entities': meaningful_entities,  # Real entity names  
                            'extracted_relationships': meaningful_relationships  # Real relationship descriptions
                        }
                        
                        logger.info(f"✅ PDF ENHANCED: Pipeline extracted {entities_count} entities and {relationships_count} relationships")
                        logger.info(f"✅ PDF ENHANCED: Created {len(semantic_concepts)} MEANINGFUL semantic concepts: {semantic_concepts[:5]}")
                        
                        # Use enhanced metadata for better concept extraction with REAL entity names
                        enhanced_content_data = type('PipelineResult', (), {
                            'concepts': [{'concept': concept} for concept in semantic_concepts[:10]],
                            'topics': meaningful_entities[:5],
                            'summary': f"PDF document containing entities: {', '.join(meaningful_entities[:5])}..." if meaningful_entities else "PDF document processed with no entities",
                            'metadata': enhanced_metadata
                        })()
                        logger.info(f"✅ PDF ENHANCED: Created enhanced content data with {len(semantic_concepts)} MEANINGFUL concepts")
                    else:
                        logger.warning(f"🔍 PDF FALLBACK: Knowledge pipeline service not available (service: {knowledge_pipeline_service is not None}, initialized: {knowledge_pipeline_service.initialized if knowledge_pipeline_service else False})")
                        enhanced_content_data = None
                        
                except Exception as e:
                    logger.error(f"❌ PDF ERROR: Error in pipeline processing: {e}")
                    import traceback
                    logger.error(f"❌ PDF ERROR: Traceback: {traceback.format_exc()}")

                    enhanced_content_data = None
                
                content = raw_content
                
            elif request.file_type == "docx":
                if not HAS_DOCX:
                    raise ValueError("DOCX processing not available - install python-docx")
                logger.info(f"Extracting text from DOCX: {file_path}")
                content = await extract_text_from_docx(file_path)
                
            else:
                # Handle text-based files
                try:
                    async with aiofiles.open(file_path, 'r', encoding=request.encoding) as f:
                        content = await f.read()
                except UnicodeDecodeError:
                    # Try with different encoding for binary files
                    try:
                        async with aiofiles.open(file_path, 'r', encoding='latin-1') as f:
                            content = await f.read()
                    except Exception:
                        async with aiofiles.open(file_path, 'rb') as f:
                            raw_content = await f.read()
                            content = f"Binary file content: {len(raw_content)} bytes - unable to extract text"
            
            progress.current_step = "Processing file content"
            progress.progress_percentage = 50.0
            progress.completed_steps = 2
            
            # Broadcast progress update
            await self._broadcast_progress_update(import_id, progress)
            
            title = request.filename
            
            logger.info(f"🔍 DEBUG: About to call _process_content for file import {import_id}")
            try:
                logger.info(f"🔍 DEBUG: Starting _process_content for text import {import_id} with timeout {CONTENT_PROCESS_TIMEOUT}s")
                processed_data = await asyncio.wait_for(
                    self._process_content(
                        content=content,
                        title=title,
                        metadata=request.source.metadata
                    ),
                    timeout=CONTENT_PROCESS_TIMEOUT
                )
                logger.info(f"🔍 DEBUG: _process_content completed for text import {import_id}")
            except AsyncioTimeoutError:
                logger.error(f"❌ Timeout during content processing for text import {import_id}")
                progress.status = "failed"
                progress.error_message = "Content processing timed out"
                await self._broadcast_progress_update(import_id, progress)
                await self._broadcast_completion(import_id, False, progress.error_message)
                return
            except Exception as e:
                logger.error(f"❌ Exception during content processing for text import {import_id}: {e}")
                logger.error(f"🔍 DEBUG: Traceback: {traceback.format_exc()}")
                progress.status = "failed"
                progress.error_message = str(e)
                await self._broadcast_progress_update(import_id, progress)
                await self._broadcast_completion(import_id, False, progress.error_message)
                return
            logger.info(f"🔍 DEBUG: _process_content returned for file import {import_id}, keys: {list(processed_data.keys())}")
            
            progress.current_step = "Creating knowledge item"
            progress.progress_percentage = 75.0
            progress.completed_steps = 3
            
            # Broadcast progress update
            await self._broadcast_progress_update(import_id, progress)
            
            # Create knowledge item with enhanced semantic data if available
            enhanced_metadata = processed_data.get('metadata', {})
            enhanced_categories = list(request.categorization_hints or ["file"])
            
            # Add enhanced semantic processing results to metadata and categories
            if enhanced_content_data:
                pipeline_metadata = getattr(enhanced_content_data, 'metadata', {})
                enhanced_metadata.update({
                    'semantic_entities': len(pipeline_metadata.get('extracted_entities', [])),
                    'semantic_relationships': len(pipeline_metadata.get('extracted_relationships', [])),
                    'pipeline_processing_time': pipeline_metadata.get('pipeline_processing_time', 0),
                    'concepts': pipeline_metadata.get('semantic_concepts', []),  # For graph extraction
                    'keywords': pipeline_metadata.get('extracted_entities', [])[:10],  # Top entities as keywords
                    'semantic_summary': getattr(enhanced_content_data, 'summary', ''),
                    'semantic_quality_score': 0.9 if pipeline_metadata.get('extracted_entities') else 0.7
                })
                
                # Add semantic topics as categories
                semantic_topics = getattr(enhanced_content_data, 'topics', [])
                if semantic_topics:
                    enhanced_categories.extend(semantic_topics[:5])  # Limit to top 5 topics
                
                logger.info(f"🔍 PDF SEMANTIC: Added semantic metadata with {len(pipeline_metadata.get('semantic_concepts', []))} concepts")
            
            knowledge_item = KnowledgeItem(
                id=f"file-{import_id}",
                content=processed_data['content'],
                knowledge_type="fact",
                title=processed_data['title'],
                source=request.source,
                import_id=import_id,
                confidence=0.8,
                quality_score=enhanced_metadata.get('semantic_quality_score', 0.8),
                categories=enhanced_categories,
                auto_categories=getattr(enhanced_content_data, 'topics', [])[:3] if enhanced_content_data else [],
                manual_categories=request.categorization_hints or ["file"],
                relationships=[],
                metadata={
                    **enhanced_metadata,
                    'filename': request.filename,
                    'file_type': request.file_type,
                    'encoding': request.encoding
                }
            )
            
            # Store the knowledge item
            await self._store_knowledge_item(knowledge_item)
            self.knowledge_store[knowledge_item.id] = knowledge_item
            
            logger.info(f"File import processed successfully: {import_id}")
            
        except Exception as e:
            logger.error(f"Failed to process file import {import_id}: {e}")
            raise

    async def _process_wikipedia_import(self, import_id: str, request):
        """Process Wikipedia import request with real API calls and progress tracking."""
        try:
            logger.info(f"🔍 DEBUG: Starting Wikipedia processing for {import_id}")
            progress = self.active_imports[import_id]
            progress.current_step = "Connecting to Wikipedia API"
            progress.progress_percentage = 10.0
            progress.completed_steps = 1
            
            # Broadcast progress update
            logger.info(f"🔍 DEBUG: Broadcasting progress update 1 for {import_id}")
            await self._broadcast_progress_update(import_id, progress)
            
            # Get Wikipedia page title from request
            page_title = getattr(request, 'page_title', getattr(request, 'topic', getattr(request, 'title', 'Unknown')))
            language = getattr(request, 'language', 'en')
            
            progress.current_step = f"Fetching Wikipedia content for '{page_title}'"
            progress.progress_percentage = 25.0
            progress.completed_steps = 2
            await self._broadcast_progress_update(import_id, progress)
            
            # Use the real Wikipedia API
            logger.info(f"🔍 DEBUG: Fetching content for page: {page_title}")
            wikipedia_data = await wikipedia_api.get_page_content(page_title, language)
            
            if wikipedia_data.get('is_fallback'):
                logger.warning(f"Using fallback content for {page_title}")
            
            content = wikipedia_data.get('content', '')
            title = wikipedia_data.get('title', page_title)
            
            # Update progress with content size information
            word_count = wikipedia_data.get('word_count', len(content.split()))
            progress.current_step = f"Processing {word_count} words from Wikipedia article"
            progress.progress_percentage = 50.0
            progress.completed_steps = 3
            
            # Broadcast progress update
            logger.info(f"🔍 DEBUG: Broadcasting progress update 2 for {import_id}")
            await self._broadcast_progress_update(import_id, progress)
            
            # Process the actual content
            logger.info(f"🔍 DEBUG: Processing content for {import_id}")
            metadata = {
                'source_url': wikipedia_data.get('url', ''),
                'language': language,
                'sections': wikipedia_data.get('sections', []),
                'summary': wikipedia_data.get('summary', ''),
                'word_count': word_count,
                'char_count': wikipedia_data.get('char_count', len(content)),
                'is_fallback': wikipedia_data.get('is_fallback', False)
            }
            
            # Merge with any existing metadata from request
            if hasattr(request, 'source') and hasattr(request.source, 'metadata'):
                metadata.update(request.source.metadata)
            
            progress.current_step = "Processing and extracting knowledge"
            progress.progress_percentage = 70.0
            progress.completed_steps = 4
            await self._broadcast_progress_update(import_id, progress)
            
            try:
                logger.info(f"🔍 DEBUG: Starting _process_content for wikipedia import {import_id} with timeout {CONTENT_PROCESS_TIMEOUT}s")
                processed_data = await asyncio.wait_for(
                    self._process_content(
                        content=content,
                        title=title,
                        metadata=metadata
                    ),
                    timeout=CONTENT_PROCESS_TIMEOUT
                )
                logger.info(f"🔍 DEBUG: _process_content completed for wikipedia import {import_id}")
            except AsyncioTimeoutError:
                logger.error(f"❌ Timeout during content processing for wikipedia import {import_id}")
                progress.status = "failed"
                progress.error_message = "Content processing timed out"
                await self._broadcast_progress_update(import_id, progress)
                await self._broadcast_completion(import_id, False, progress.error_message)
                return
            except Exception as e:
                logger.error(f"❌ Exception during content processing for wikipedia import {import_id}: {e}")
                logger.error(f"🔍 DEBUG: Traceback: {traceback.format_exc()}")
                progress.status = "failed"
                progress.error_message = str(e)
                await self._broadcast_progress_update(import_id, progress)
                await self._broadcast_completion(import_id, False, progress.error_message)
                return
            
            progress.current_step = "Creating knowledge item"
            progress.progress_percentage = 85.0
            progress.completed_steps = 5
            
            # Broadcast progress update
            logger.info(f"🔍 DEBUG: Broadcasting progress update 3 for {import_id}")
            await self._broadcast_progress_update(import_id, progress)
            
            # Create knowledge item with real data
            logger.info(f"🔍 DEBUG: Creating knowledge item for {import_id}")
            knowledge_item = KnowledgeItem(
                id=f"wikipedia-{import_id}",
                content=processed_data['content'],
                knowledge_type="fact",
                title=processed_data['title'],
                source=ImportSource(
                    source_type="wikipedia",
                    source_identifier=page_title,
                    metadata=metadata
                ),
                import_id=import_id,
                confidence=0.9 if not wikipedia_data.get('is_fallback') else 0.3,
                quality_score=0.8,
                categories=request.categorization_hints or ["wikipedia"],
                auto_categories=[],
                manual_categories=request.categorization_hints or ["wikipedia"],
                relationships=[],
                metadata={
                    **processed_data.get('metadata', {}),
                    'wikipedia_page': request.page_title,
                    'language': request.language
                }
            )
            logger.info(f"🔍 DEBUG: Knowledge item created for {import_id}")
            
            # Store the knowledge item
            logger.info(f"🔍 DEBUG: Storing knowledge item for {import_id}")
            await self._store_knowledge_item(knowledge_item)
            self.knowledge_store[knowledge_item.id] = knowledge_item
            logger.info(f"🔍 DEBUG: Knowledge item stored for {import_id}")
            
            progress.current_step = "Wikipedia import completed"
            progress.progress_percentage = 100.0
            progress.completed_steps = 4
            
            # Broadcast final progress update
            logger.info(f"🔍 DEBUG: Broadcasting final progress update for {import_id}")
            await self._broadcast_progress_update(import_id, progress)
            
            logger.info(f"Wikipedia import processed successfully: {import_id}")
            
        except Exception as e:
            logger.error(f"Failed to process Wikipedia import {import_id}: {e}")
            raise

    async def _store_knowledge_item(self, knowledge_item: KnowledgeItem):
        """Store a knowledge item to persistent storage and knowledge management service."""
        try:
            # Store to local file system
            file_path = self.storage_path / f"{knowledge_item.id}.json"
            file_write_start = time.perf_counter()
            async with aiofiles.open(file_path, 'w') as f:
                await f.write(knowledge_item.model_dump_json(indent=2))
            file_write_dur = time.perf_counter() - file_write_start
            logger.debug(f"Stored knowledge item: {knowledge_item.id} (file write {file_write_dur:.3f}s)")
            
            # UNIFIED KNOWLEDGE GRAPH: Use transparency knowledge graph as the single source of truth
            # This eliminates the dual graph anti-pattern and ensures consistency
            tkg_start = time.perf_counter()
            await self._add_to_transparency_knowledge_graph(knowledge_item)
            tkg_dur = time.perf_counter() - tkg_start
            logger.info(f"🔍 UNIFIED GRAPH: Added item {knowledge_item.id} to unified knowledge graph in {tkg_dur:.3f}s")
            
            # TODO: Remove knowledge_management_service dependency entirely - it's redundant and causes inconsistency
            # Legacy code attempted to maintain two separate graph systems which is a bad design pattern

            # Broadcast knowledge update via WebSocket (measure duration)
            if self.websocket_manager and self.websocket_manager.has_connections():
                try:
                    # Ensure the payload is JSON serializable (convert Pydantic models to dicts)
                    try:
                        if hasattr(knowledge_item.source, "model_dump"):
                            source_serializable = knowledge_item.source.model_dump()
                        elif hasattr(knowledge_item.source, "dict"):
                            source_serializable = knowledge_item.source.dict()
                        else:
                            source_serializable = str(knowledge_item.source)
                    except Exception:
                        source_serializable = str(knowledge_item.source)

                    broadcast_start = time.perf_counter()
                    # Get current document count by counting stored knowledge items
                    document_count = len([f for f in self.storage_path.glob("*.json") if not f.name.startswith("temp_")])
                    
                    await self.websocket_manager.broadcast({
                        "type": "knowledge_update",
                        "event": "item_added",
                        "data": {
                            "item_id": knowledge_item.id,
                            "title": knowledge_item.title,
                            "source": source_serializable,
                            "categories": knowledge_item.categories,
                            "timestamp": time.time()
                        },
                        "stats": {
                            "totalDocuments": document_count,
                            "newDocument": True,
                            "documentType": source_serializable.get("source_type", "unknown") if isinstance(source_serializable, dict) else "unknown"
                        }
                    })
                    broadcast_dur = time.perf_counter() - broadcast_start
                    logger.info(f"🔍 KNOWLEDGE BROADCAST: Broadcasted knowledge update for {knowledge_item.id} in {broadcast_dur:.3f}s")
                except Exception as e:
                    logger.error(f"Failed to broadcast knowledge update for {knowledge_item.id}: {e}")
                    
        except Exception as e:
            logger.error(f"Failed to store knowledge item {knowledge_item.id}: {e}")
            raise


    async def _add_to_transparency_knowledge_graph(self, knowledge_item: KnowledgeItem):
        """Add knowledge item to the cognitive transparency knowledge graph for visualization."""
        try:
            # Import here to avoid circular dependency
            from backend.cognitive_transparency_integration import cognitive_transparency_api
            
            logger.info(f"🔍 GRAPH SYNC: Starting to add knowledge item {knowledge_item.id} to transparency graph")
            logger.info(f"🔍 GRAPH SYNC: cognitive_transparency_api exists: {cognitive_transparency_api is not None}")
            logger.info(f"🔍 GRAPH SYNC: knowledge_graph exists: {cognitive_transparency_api.knowledge_graph is not None if cognitive_transparency_api else False}")
            
            # FIXED: Don't create fallback instances - this was causing dual graph problem!
            # Instead, wait for proper initialization or skip if not ready
            if not cognitive_transparency_api or not getattr(cognitive_transparency_api, 'knowledge_graph', None):
                logger.warning(f"🔍 GRAPH SYNC: Transparency API not fully initialized yet for item {knowledge_item.id}, skipping graph sync")
                return  # Skip graph sync if not properly initialized

            if cognitive_transparency_api and cognitive_transparency_api.knowledge_graph:
                # Extract MEANINGFUL concepts from the knowledge item for graph nodes
                concepts = []
                
                # Define the same filtering function for graph concepts
                def is_meaningful_graph_concept(text: str) -> bool:
                    """Filter for semantically meaningful graph concepts only."""
                    if not text or len(text.strip()) < 3:
                        return False
                    
                    text = text.strip()
                    
                    # Skip generic/noise terms
                    noise_terms = {
                        'file', 'document', 'pdf', 'docx', 'txt', 'upload', 'download',
                        'the', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for', 'of', 'with',
                        'this', 'that', 'these', 'those', 'what', 'where', 'when', 'how', 'why',
                        'today', 'yesterday', 'tomorrow', 'now', 'then', 'here', 'there',
                        'manual', 'web', 'wikipedia'  # Skip generic categories
                    }
                    if text.lower() in noise_terms:
                        return False
                    
                    # Skip pure numbers or single characters
                    if text.isdigit() or len(text) == 1:
                        return False
                    
                    # Skip file extensions and filename patterns
                    if ('.' in text and len(text.split('.')[-1]) <= 4) or text.endswith('.pdf'):
                        return False
                    
                    # Include multi-word technical terms (likely meaningful)
                    if ' ' in text and len(text) > 6:
                        return True
                    
                    # Include capitalized terms (likely proper nouns) but not all-caps single words
                    if text[0].isupper() and len(text) > 3 and not text.isupper():
                        return True
                    
                    # Include terms with mixed case (likely technical/compound terms)
                    if any(c.isupper() for c in text[1:]) and len(text) > 4:
                        return True
                    
                    return False
                
                # PRIORITIZE: Add semantic concepts from pipeline processing FIRST (most meaningful)
                if knowledge_item.metadata and 'concepts' in knowledge_item.metadata:
                    semantic_concepts = knowledge_item.metadata['concepts']
                    if isinstance(semantic_concepts, list):
                        filtered_semantic = [c for c in semantic_concepts if is_meaningful_graph_concept(c)]
                        concepts.extend(filtered_semantic[:6])  # Top 6 semantic concepts
                        logger.info(f"🔍 GRAPH SYNC: Added FILTERED semantic pipeline concepts: {filtered_semantic[:6]}")
                
                # Add semantic entity keywords (high value entities from NLP)
                if knowledge_item.metadata and 'keywords' in knowledge_item.metadata:
                    keywords = knowledge_item.metadata['keywords']
                    if isinstance(keywords, list):
                        filtered_keywords = [k for k in keywords if is_meaningful_graph_concept(k)]
                        concepts.extend(filtered_keywords[:4])  # Top 4 entity keywords
                        logger.info(f"🔍 GRAPH SYNC: Added FILTERED semantic entity keywords: {filtered_keywords[:4]}")
                
                # Only add title if it's meaningful (not just a filename)
                if knowledge_item.title and is_meaningful_graph_concept(knowledge_item.title):
                    concepts.append(knowledge_item.title)
                    logger.info(f"🔍 GRAPH SYNC: Added FILTERED title concept: {knowledge_item.title}")
                else:
                    logger.info(f"🔍 GRAPH SYNC: Skipped non-meaningful title: {knowledge_item.title}")
                
                # SKIP generic categories entirely - they add no semantic value
                logger.info(f"🔍 GRAPH SYNC: Skipping generic categories to avoid noise: {knowledge_item.categories}")
                
                # ENHANCED: Add semantic topics from pipeline processing  
                if knowledge_item.metadata and 'semantic_summary' in knowledge_item.metadata:
                    summary = knowledge_item.metadata['semantic_summary']
                    if summary and len(summary) > 10:
                        # Extract key terms from semantic summary
                        import re
                        key_terms = re.findall(r'\b[A-Z][a-z]+\b', summary)
                        filtered_terms = [t for t in key_terms if is_meaningful_graph_concept(t)]
                        if filtered_terms:
                            concepts.extend(filtered_terms[:2])  # Top 2 summary concepts
                            logger.info(f"🔍 GRAPH SYNC: Added FILTERED semantic summary concepts: {filtered_terms[:2]}")
                
                # Remove duplicates while preserving order
                unique_concepts = []
                seen = set()
                for concept in concepts:
                    if concept not in seen:
                        unique_concepts.append(concept)
                        seen.add(concept)
                concepts = unique_concepts
                
                logger.info(f"🔍 GRAPH SYNC: Total MEANINGFUL concepts to add: {len(concepts)} - {concepts}")
                
                # Add each concept as a node in the knowledge graph
                for concept in concepts:
                    if concept and isinstance(concept, str) and len(concept.strip()) > 0:
                        try:
                            logger.info(f"🔍 GRAPH SYNC: Attempting to add concept '{concept}' to knowledge graph")
                            result = cognitive_transparency_api.knowledge_graph.add_node(
                                concept=concept.strip(),
                                node_type="knowledge_item",
                                properties={
                                    "source_item_id": knowledge_item.id,
                                    "source": knowledge_item.source.source_type if knowledge_item.source else "unknown",
                                    "confidence": knowledge_item.confidence,
                                    "quality_score": knowledge_item.quality_score
                                },
                                confidence=knowledge_item.confidence
                            )
                            logger.info(f"🔍 GRAPH SYNC: Successfully added concept '{concept}' to knowledge graph, result: {result}")
                        except Exception as e:
                            logger.warning(f"🔍 GRAPH SYNC: Failed to add concept '{concept}' to knowledge graph: {e}")
                            logger.warning(f"🔍 GRAPH SYNC: Exception details: {type(e).__name__}: {str(e)}")
                
                # Create relationships between concepts from the same item
                if len(concepts) > 1:
                    main_concept = concepts[0]  # Use title as main concept
                    for related_concept in concepts[1:]:
                        if related_concept and isinstance(related_concept, str) and len(related_concept.strip()) > 0:
                            try:
                                logger.info(f"🔍 GRAPH SYNC: Attempting to add relationship '{main_concept}' -> '{related_concept}'")
                                result = cognitive_transparency_api.knowledge_graph.add_edge(
                                    source_concept=main_concept.strip(),
                                    target_concept=related_concept.strip(),
                                    relation_type="related_to",
                                    strength=0.7,
                                    properties={
                                        "source_item_id": knowledge_item.id,
                                        "relationship_source": "knowledge_ingestion"
                                    },
                                    confidence=0.7
                                )
                                logger.info(f"🔍 GRAPH SYNC: Successfully added relationship '{main_concept}' -> '{related_concept}', result: {result}")
                            except Exception as e:
                                logger.warning(f"🔍 GRAPH SYNC: Failed to add relationship '{main_concept}' -> '{related_concept}': {e}")
                
                # Broadcast knowledge graph update to frontend
                if self.websocket_manager and self.websocket_manager.has_connections():
                    try:
                        # Export updated graph data and ensure it's serializable
                        graph_data = await cognitive_transparency_api.knowledge_graph.export_graph()
                        logger.info(f"🔍 GRAPH SYNC: Exported graph data has {len(graph_data.get('nodes', []))} nodes and {len(graph_data.get('edges', []))} edges")
                        
                        try:
                            nodes = graph_data.get("nodes", [])
                            links = graph_data.get("edges", [])

                            # Serialize nodes and links defensively
                            serializable_nodes = []
                            for n in nodes:
                                try:
                                    if hasattr(n, "model_dump"):
                                        serializable_nodes.append(n.model_dump())
                                    elif hasattr(n, "dict"):
                                        serializable_nodes.append(n.dict())
                                    else:
                                        serializable_nodes.append(dict(n) if isinstance(n, (list, tuple)) else n)
                                except Exception:
                                    serializable_nodes.append(str(n))

                            serializable_links = []
                            for l in links:
                                try:
                                    if hasattr(l, "model_dump"):
                                        serializable_links.append(l.model_dump())
                                    elif hasattr(l, "dict"):
                                        serializable_links.append(l.dict())
                                    else:
                                        serializable_links.append(dict(l) if isinstance(l, (list, tuple)) else l)
                                except Exception:
                                    serializable_links.append(str(l))

                            await self.websocket_manager.broadcast({
                                "type": "knowledge-graph-update",
                                "data": {
                                    "nodes": serializable_nodes,
                                    "links": serializable_links,
                                    "timestamp": time.time(),
                                    "update_source": "knowledge_ingestion"
                                }
                            })
                            logger.info(f"🔍 GRAPH SYNC: Broadcasted updated knowledge graph with {len(serializable_nodes)} nodes")
                        except Exception as e:
                            logger.warning(f"🔍 GRAPH SYNC: Failed to serialize or broadcast graph data: {e}")
                    except Exception as e:
                        logger.warning(f"🔍 GRAPH SYNC: Failed to broadcast knowledge graph update: {e}")
                        
            else:
                logger.warning(f"🔍 GRAPH SYNC: Cognitive transparency knowledge graph not available")
                
        except Exception as e:
            logger.error(f"🔍 GRAPH SYNC: Failed to add knowledge item {knowledge_item.id} to transparency knowledge graph: {e}")
            logger.error(f"🔍 GRAPH SYNC: Exception details: {type(e).__name__}: {str(e)}")
            # Don't raise the exception as this is not critical for the ingestion process
    
    async def _send_wikipedia_progress_updates(self, import_id: str, request: WikipediaImportRequest):
        """Send WebSocket progress updates for Wikipedia import processing."""
        logger.info(f"🔍 BACKGROUND TASK: Started for {import_id}")
        try:
            # Wait a tiny bit for the initial request to return
            await asyncio.sleep(0.1)
            logger.info(f"🔍 BACKGROUND TASK: Ready to send updates for {import_id}")
            
            # Progress update 1: Starting
            if self.websocket_manager and self.websocket_manager.has_connections():
                progress_event = {
                    "type": "import_progress",
                    "timestamp": time.time(),
                    "import_id": import_id,
                    "progress": 25.0,
                    "status": "processing",
                    "message": "Fetching Wikipedia content...",
                    "completed_steps": 1,
                    "total_steps": 4
                }
                await self.websocket_manager.broadcast(progress_event)
                logger.info(f"🔍 PROGRESS UPDATE: Sent 25% for {import_id}")
            
            await asyncio.sleep(0.5)
            
            # Progress update 2: Processing
            if self.websocket_manager and self.websocket_manager.has_connections():
                progress_event = {
                    "type": "import_progress",
                    "timestamp": time.time(),
                    "import_id": import_id,
                    "progress": 50.0,
                    "status": "processing",
                    "message": "Processing Wikipedia content...",
                    "completed_steps": 2,
                    "total_steps": 4
                }
                await self.websocket_manager.broadcast(progress_event)
                logger.info(f"🔍 PROGRESS UPDATE: Sent 50% for {import_id}")
            
            await asyncio.sleep(0.5)
            
            # Progress update 3: Creating
            if self.websocket_manager and self.websocket_manager.has_connections():
                progress_event = {
                    "type": "import_progress",
                    "timestamp": time.time(),
                    "import_id": import_id,
                    "progress": 75.0,
                    "status": "processing",
                    "message": "Creating knowledge item...",
                    "completed_steps": 3,
                    "total_steps": 4
                }
                await self.websocket_manager.broadcast(progress_event)
                logger.info(f"🔍 PROGRESS UPDATE: Sent 75% for {import_id}")
            
            # Wait for actual processing to complete (look for completion in active_imports)
            max_wait = 10  # 10 seconds max
            wait_time = 0
            while wait_time < max_wait:
                await asyncio.sleep(0.2)
                wait_time += 0.2
                
                if import_id in self.active_imports:
                    progress = self.active_imports[import_id]
                    if progress.status == "completed":
                        # Send completion event
                        if self.websocket_manager and self.websocket_manager.has_connections():
                            completion_event = {
                                "type": "import_completed",
                                "timestamp": time.time(),
                                "import_id": import_id,
                                "message": "Wikipedia import completed successfully",
                                "success": True
                            }
                            await self.websocket_manager.broadcast(completion_event)
                            logger.info(f"🔍 COMPLETION: Sent completion for {import_id}")
                        break
                    elif progress.status == "failed":
                        # Send failure event
                        if self.websocket_manager and self.websocket_manager.has_connections():
                            failure_event = {
                                "type": "import_failed",
                                "timestamp": time.time(),
                                "import_id": import_id,
                                "message": f"Wikipedia import failed: {progress.error_message}",
                                "success": False
                            }
                            await self.websocket_manager.broadcast(failure_event)
                            logger.info(f"🔍 FAILURE: Sent failure for {import_id}")
                        break
            
        except Exception as e:
            logger.error(f"🔍 ERROR: Failed to send progress updates for {import_id}: {e}")


# Global instance
knowledge_ingestion_service = KnowledgeIngestionService()
